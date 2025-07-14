# the below program generates checklist docx format 
# pip install python-docx

from docx import Document
from docx.shared import Pt

# Create the document
doc = Document()
doc.add_heading('AP Computer Science Principles (CSP) Topic Checklist', level=1)

# Define Big Ideas and their topics
topics = {
    "Big Idea 1: Creative Development": [
        "Program development lifecycle (define, design, implement, test)",
        "Algorithms and problem-solving",
        "Documentation and debugging",
        "Code readability and clarity",
        "Pair programming & collaboration tools"
    ],
    "Big Idea 2: Data": [
        "Binary representation of data (bits and bytes)",
        "Text, images, and numbers in binary",
        "Data compression (lossy vs. lossless)",
        "Data storage and file sizes",
        "Analyzing large data sets",
        "Visualizing trends (charts, graphs)",
        "Bias and ethical concerns in data use"
    ],
    "Big Idea 3: Computing Systems and Networks": [
        "Structure of the internet (clients, servers, routers)",
        "IP addresses, DNS, and packet-switching",
        "Fault tolerance and redundancy",
        "Encryption and secure data transmission",
        "Phishing, malware, and cyber threats",
        "Protecting personal data"
    ],
    "Big Idea 4: Algorithms and Programming": [
        "Variables and data types",
        "Arithmetic and logical operations",
        "Input and output",
        "Conditionals (if, else)",
        "Loops (for, while)",
        "Functions (parameters, return values)",
        "Lists/arrays and list operations (append, insert, access)",
        "Traversing lists",
        "Developing, testing, and refining programs",
        "Abstraction through procedures and functions"
    ],
    "Big Idea 5: Impact of Computing": [
        "Positive and negative impacts of computing",
        "Digital divide and global access to technology",
        "Intellectual property, copyright, fair use",
        "Privacy and data ethics",
        "Responsible computing (accessibility, inclusivity)"
    ],
    "Performance Task (Create Task)": [
        "Design and build a program",
        "Include input, list, and algorithm (sequence, selection, iteration)",
        "Submit code and written reflection"
    ]
}

# Add topics to the document
for big_idea, subtopics in topics.items():
    doc.add_heading(big_idea, level=2)
    for topic in subtopics:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(topic)
        run.font.size = Pt(11)

# Save the document
doc.save("AP_CSP_Topic_Checklist.docx")

